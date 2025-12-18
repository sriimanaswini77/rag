# -------------------------------------------------
# IMPORTS
# -------------------------------------------------
import os, json, math, requests
import numpy as np
from dotenv import load_dotenv
from datasets import Dataset

# LangChain
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_ollama import OllamaEmbeddings, ChatOllama

# RAGAS
from ragas import evaluate
from ragas.metrics import context_precision, context_recall, answer_relevancy, faithfulness

# PDF + DOCX
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from docx import Document

# Generation Metrics
import nltk
from nltk.translate.bleu_score import sentence_bleu, SmoothingFunction
from rouge_score import rouge_scorer
from nltk.translate.meteor_score import meteor_score
from bert_score import score as bert_score

nltk.download('punkt')

load_dotenv()


# -------------------------------------------------
# DOCUMENT
# -------------------------------------------------
markdown = """Deep learning is a powerful subset of machine learning within AI using multi-layer neural networks..."""


# -------------------------------------------------
# CHUNKING
# -------------------------------------------------
splitter = RecursiveCharacterTextSplitter(chunk_size=80, chunk_overlap=10)
chunks = splitter.create_documents([markdown])

for i, c in enumerate(chunks):
    c.metadata["id"] = str(i)


# -------------------------------------------------
# FAISS VECTOR STORE
# -------------------------------------------------
embeddings = OllamaEmbeddings(model="qwen3-embedding:0.6b")
vector_store = FAISS.from_documents(chunks, embeddings)
retriever = vector_store.as_retriever(search_kwargs={"k": 5})


# -------------------------------------------------
# LLM
# -------------------------------------------------
llm = ChatOllama(model="mistral-small-3.1-24b:latest", temperature=0.2)


# -------------------------------------------------
# PROMPT
# -------------------------------------------------
prompt = PromptTemplate(
    template="""
Answer ONLY from the provided context. If not found, say "I don't know."

Context:
{context}

Question: {question}
""",
    input_variables=["context", "question"]
)


# -------------------------------------------------
# QUESTIONS
# -------------------------------------------------
questions = [
    "What is deep learning and tell me its applications?",
    "How does a neural network learn hierarchical features?"
]

references = [
    "Deep learning is a subset of ML using neural networks.",
    "A neural network learns hierarchical features through layered processing."
]

answers, contexts, retrieved_ids_per_q = [], [], []


# -------------------------------------------------
# RAG PIPELINE: RETRIEVAL + GENERATION
# -------------------------------------------------
for q in questions:
    retrieved_docs = retriever.invoke(q)
    retrieved_ids_per_q.append([d.metadata["id"] for d in retrieved_docs])

    ctx = "\n\n".join(d.page_content for d in retrieved_docs)
    final_prompt = prompt.invoke({"context": ctx, "question": q})
    answer = llm.invoke(final_prompt).content

    answers.append(answer)
    contexts.append([d.page_content for d in retrieved_docs])


# -------------------------------------------------
# IR METRICS
# -------------------------------------------------
def hit_rate(y_true, y_pred):
    return sum(any(doc in t for doc in p) for t, p in zip(y_true, y_pred)) / len(y_true)

def mean_reciprocal_rank(y_true, y_pred):
    out = []
    for t, p in zip(y_true, y_pred):
        rr = 0
        for rank, doc in enumerate(p, start=1):
            if doc in t:
                rr = 1/rank
                break
        out.append(rr)
    return float(np.mean(out))

def ndcg(y_true, y_pred, k=5):
    out = []
    for t, p in zip(y_true, y_pred):
        pk = p[:k]
        dcg = sum(1/np.log2(i+2) if doc in t else 0 for i, doc in enumerate(pk))
        ideal = t[:k]
        idcg = sum(1/np.log2(i+2) for i in range(len(ideal)))
        out.append(dcg/idcg if idcg else 0)
    return float(np.mean(out))

def recall_precision_at_k(y_true, y_pred, k=5):
    recall, precision = [], []
    for t, p in zip(y_true, y_pred):
        pk = p[:k]
        hits = sum(doc in t for doc in pk)
        recall.append(hits/len(t))
        precision.append(hits/k)
    return {"Recall": np.mean(recall), "Precision": np.mean(precision)}

relevant_ids = [c.metadata["id"] for c in chunks if "deep learning" in c.page_content.lower()]
y_true = [relevant_ids for _ in questions]
y_pred = retrieved_ids_per_q

ir_metrics_per_q = []
k = 5

for i in range(len(questions)):
    rp = recall_precision_at_k([y_true[i]], [y_pred[i]], k)
    ir_metrics_per_q.append({
        "question": questions[i],
        "precision@k": rp["Precision"],
        "recall@k": rp["Recall"],
        "hit_rate@k": hit_rate([y_true[i]], [y_pred[i]]),
        "MRR@k": mean_reciprocal_rank([y_true[i]], [y_pred[i]]),
        "NDCG@k": ndcg([y_true[i]], [y_pred[i]], k),
    })


# -------------------------------------------------
# GENERATION METRICS (ONLY REQUESTED ONES)
# -------------------------------------------------

# 1. BLEU
def compute_bleu(preds, refs):
    smooth = SmoothingFunction().method1
    scores = [sentence_bleu([nltk.word_tokenize(r)],
                            nltk.word_tokenize(p),
                            smoothing_function=smooth)
              for p, r in zip(preds, refs)]
    return float(np.mean(scores))

# 2. ROUGE
def compute_rouge(preds, refs):
    scorer = rouge_scorer.RougeScorer(["rouge1", "rougeL"], use_stemmer=True)
    r1, rL = [], []
    for p, r in zip(preds, refs):
        s = scorer.score(r, p)
        r1.append(s["rouge1"].fmeasure)
        rL.append(s["rougeL"].fmeasure)
    return float(np.mean(r1)), float(np.mean(rL))

# 3. METEOR
def compute_meteor(preds, refs):
    return float(np.mean([meteor_score([r], p) for p, r in zip(preds, refs)]))

# 4. BERTScore
def compute_bertscore(preds, refs):
    _, _, F1 = bert_score(preds, refs, lang="en", rescale_with_baseline=True)
    return float(F1.mean())

# 5. Perplexity (Ollama logprobs)
def ollama_perplexity(text, model="mistral-small-3.1-24b:latest"):
    url = "http://localhost:11434/api/generate"
    payload = {"model": model, "prompt": text, "logprobs": 1, "stream": False}
    res = requests.post(url, json=payload).json()

    if "logprobs" not in res: return None
    lp = [tok["logprob"] for tok in res["logprobs"]["tokens"]]
    if len(lp) == 0: return None

    avg_lp = sum(lp)/len(lp)
    return math.exp(-avg_lp)


# Compute generation metrics
bleu = compute_bleu(answers, references)
rouge1, rougeL = compute_rouge(answers, references)
meteor = compute_meteor(answers, references)
bertscore_f1 = compute_bertscore(answers, references)
perplexities = [ollama_perplexity(a) for a in answers]


# -------------------------------------------------
# RAGAS METRICS
# -------------------------------------------------
ragas_results_per_q = []
for i in range(len(questions)):
    ds = Dataset.from_dict({
        "question": [questions[i]],
        "answer": [answers[i]],
        "contexts": [contexts[i]],
        "reference": [references[i]],
    })

    result = evaluate(ds, metrics=[
        context_precision, context_recall, answer_relevancy, faithfulness
    ], llm=llm, embeddings=embeddings)

    try:
        ragas_dict = result.to_dict()
    except:
        ragas_dict = result.to_pandas().to_dict(orient="records")[0]

    ragas_results_per_q.append({"question": questions[i], "metrics": ragas_dict})


# -------------------------------------------------
# SAVE DOCX
# -------------------------------------------------
doc = Document()
doc.add_heading("RAG Evaluation Report", level=1)

# Generation Metrics (Global)
doc.add_heading("Generation Metrics", level=2)
doc.add_paragraph(f"BLEU: {bleu:.4f}")
doc.add_paragraph(f"ROUGE-1: {rouge1:.4f}")
doc.add_paragraph(f"ROUGE-L: {rougeL:.4f}")
doc.add_paragraph(f"METEOR: {meteor:.4f}")
doc.add_paragraph(f"BERTScore F1: {bertscore_f1:.4f}")
doc.add_paragraph(f"Perplexities: {perplexities}")

for i in range(len(questions)):
    q = questions[i]
    ir = ir_metrics_per_q[i]
    ragas = ragas_results_per_q[i]["metrics"]

    doc.add_heading(q, level=2)
    doc.add_paragraph("IR Metrics:")
    for k, v in ir.items():
        if k != "question": doc.add_paragraph(f"{k}: {v:.4f}", style="List Bullet")

    doc.add_paragraph("RAGAS Metrics:")
    for k, v in ragas.items():
        doc.add_paragraph(f"{k}: {v:.4f}", style="List Bullet")

doc.save("rag_evaluation_report.docx")


# -------------------------------------------------
# SAVE PDF
# -------------------------------------------------
styles = getSampleStyleSheet()
small = ParagraphStyle("small", parent=styles["Normal"], fontSize=8)
title = ParagraphStyle("title", parent=styles["Title"], fontSize=11)

flow = []
pdf = SimpleDocTemplate("rag_evaluation_report.pdf", pagesize=A4,
                        topMargin=10, bottomMargin=10, leftMargin=10, rightMargin=10)

flow.append(Paragraph("<b><i>Generation Metrics</i></b>", title))
flow.append(Paragraph(f"BLEU: {bleu:.4f}", small))
flow.append(Paragraph(f"ROUGE-1: {rouge1:.4f}", small))
flow.append(Paragraph(f"ROUGE-L: {rougeL:.4f}", small))
flow.append(Paragraph(f"METEOR: {meteor:.4f}", small))
flow.append(Paragraph(f"BERTScore F1: {bertscore_f1:.4f}", small))
flow.append(Paragraph(f"Perplexities: {perplexities}", small))
flow.append(Spacer(1, 8))

for i in range(len(questions)):
    q = questions[i]
    ir = ir_metrics_per_q[i]
    ragas = ragas_results_per_q[i]["metrics"]

    flow.append(Paragraph(f"<b><i>Question: {q}</i></b>", title))

    ir_rows = [["Metric", "Value"]] + [[k, f"{v:.4f}"] for k, v in ir.items() if k != "question"]
    ir_table = Table(ir_rows, colWidths=[60*mm, 25*mm])
    ir_table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
                                  ("GRID", (0,0), (-1,-1), 0.25, colors.grey)]))
    flow.append(ir_table)

    ragas_rows = [["Metric", "Value"]] + [[k, f"{v:.4f}"] for k, v in ragas.items()]
    ragas_table = Table(ragas_rows, colWidths=[60*mm, 25*mm])
    ragas_table.setStyle(TableStyle([("BACKGROUND", (0,0), (-1,0), colors.lightgrey),
                                     ("GRID", (0,0), (-1,-1), 0.25, colors.grey)]))

    flow.append(ragas_table)
    flow.append(Spacer(1, 8))

pdf.build(flow)


print("\nReport Saved Successfully!")
print("PDF  -> rag_evaluation_report.pdf")
print("DOCX -> rag_evaluation_report.docx")






########################## REQUIREMENTS #################
'''nltk
rouge-score
bert-score
torch
requests
reportlab
python-docx
'''
