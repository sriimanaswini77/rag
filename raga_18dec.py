# -------------------------------------------------
# 0️⃣  Imports – unchanged + new ones for PDF/DOCX
# -------------------------------------------------
import os, json
from pathlib import Path
from dotenv import load_dotenv
import numpy as np
from datasets import Dataset

# LangChain / Ollama
from langchain_classic.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import PromptTemplate
from langchain_ollama import OllamaEmbeddings, ChatOllama

# RAGAS
from ragas import evaluate
from ragas.metrics import (
    context_precision,
    context_recall,
    answer_relevancy,
    faithfulness,
)

# NEW IMPORTS FOR PDF + DOCX OUTPUT
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from docx import Document

load_dotenv()


# -------------------------------------------------
# 1️⃣ INPUT DOCUMENT
# -------------------------------------------------
markdown = '''Deep learning is a powerful subset of machine learning within AI that uses multi-layered artificial neural networks ...'''


# -------------------------------------------------
# 2️⃣ CHUNKING
# -------------------------------------------------
splitter = RecursiveCharacterTextSplitter(chunk_size=80, chunk_overlap=10)
chunks = splitter.create_documents([markdown])


# -------------------------------------------------
# 3️⃣ ADD METADATA IDs
# -------------------------------------------------
for i, chunk in enumerate(chunks):
    chunk.metadata["id"] = str(i)


# -------------------------------------------------
# 4️⃣ FAISS VECTOR STORE
# -------------------------------------------------
embeddings = OllamaEmbeddings(model="qwen3-embedding:0.6b")
vector_store = FAISS.from_documents(chunks, embeddings)
retriever = vector_store.as_retriever(
    search_type="similarity",
    search_kwargs={"k": 5},
)


# -------------------------------------------------
# 5️⃣ LLM
# -------------------------------------------------
llm = ChatOllama(
    model="mistral-small-3.1-24b:latest",
    temperature=0.2
)


# -------------------------------------------------
# 6️⃣ PROMPT TEMPLATE
# -------------------------------------------------
prompt = PromptTemplate(
    template="""
You are a helpful assistant.
Answer ONLY from the provided context.
If the answer isn't in the context, say "I don't know."

Context:
{context}

Question: {question}
""",
    input_variables=["context", "question"],
)


# -------------------------------------------------
# 7️⃣ MULTIPLE QUESTIONS
# -------------------------------------------------
questions = [
    "What is deep learning and tell me its applications?",
    "How does a neural network learn hierarchical features?"
]

references = [
    "Deep learning is a subset of ML using neural networks.",
    "A neural network learns hierarchical features through layers."
]

answers = []
contexts = []
retrieved_ids_per_q = []

for q in questions:
    retrieved_docs = retriever.invoke(q)
    retrieved_ids_per_q.append([doc.metadata["id"] for doc in retrieved_docs])

    context_text = "\n\n".join(doc.page_content for doc in retrieved_docs)

    final_prompt = prompt.invoke({"context": context_text, "question": q})
    answer = llm.invoke(final_prompt).content
    answers.append(answer)

    contexts.append([doc.page_content for doc in retrieved_docs])


# -------------------------------------------------
# 8️⃣ UPDATED IR METRICS (your custom functions)
# -------------------------------------------------
def hit_rate(y_true, y_pred):
    return sum(any(doc in t for doc in p) for t, p in zip(y_true, y_pred)) / len(y_true)

def mean_reciprocal_rank(y_true, y_pred):
    rrs = []
    for t, p in zip(y_true, y_pred):
        rr = 0
        for rank, doc in enumerate(p, start=1):
            if doc in t:
                rr = 1 / rank
                break
        rrs.append(rr)
    return np.mean(rrs)

def ndcg(y_true, y_pred, k=5):
    scores = []
    for t, p in zip(y_true, y_pred):
        p_k = p[:k]
        dcg = sum(1 / np.log2(i + 2) if doc in t else 0 for i, doc in enumerate(p_k))
        ideal = t[:k]
        idcg = sum(1 / np.log2(i + 2) for i in range(len(ideal)))
        scores.append(dcg / idcg if idcg else 0)
    return np.mean(scores)

def recall_precision_at_k(y_true, y_pred, k=5):
    recalls, precisions = [], []
    for t, p in zip(y_true, y_pred):
        p_k = p[:k]
        hits = sum(doc in t for doc in p_k)
        recalls.append(hits / len(t))
        precisions.append(hits / k)
    return {
        f"Recall@{k}": np.mean(recalls),
        f"Precision@{k}": np.mean(precisions)
    }


# Relevant IDs for evaluation
relevant_ids = [
    chunk.metadata["id"]
    for chunk in chunks
    if "deep learning" in chunk.page_content.lower()
]

k = 5
y_pred = retrieved_ids_per_q
y_true = [relevant_ids for _ in questions]

ir_metrics_per_q = []

for i in range(len(questions)):
    true_list = [y_true[i]]
    pred_list = [y_pred[i]]
    rp = recall_precision_at_k(true_list, pred_list, k)

    ir_metrics_per_q.append({
        "question": questions[i],
        "precision@k": rp[f"Precision@{k}"],
        "recall@k": rp[f"Recall@{k}"],
        "hit_rate@k": hit_rate(true_list, pred_list),
        "MRR@k": mean_reciprocal_rank(true_list, pred_list),
        "NDCG@k": ndcg(true_list, pred_list, k),
    })


# -------------------------------------------------
# 9️⃣ RAGAS METRICS
# -------------------------------------------------
ragas_results_per_q = []

for i in range(len(questions)):
    ragas_dataset = Dataset.from_dict({
        "question": [questions[i]],
        "answer": [answers[i]],
        "contexts": [contexts[i]],
        "reference": [references[i]],
    })

    ragas_result = evaluate(
        dataset=ragas_dataset,
        metrics=[context_precision, context_recall, answer_relevancy, faithfulness],
        llm=llm,
        embeddings=embeddings,
    )

    try:
        ragas_dict = ragas_result.to_dict()
    except:
        ragas_dict = ragas_result.to_pandas().to_dict(orient="records")[0]

    ragas_results_per_q.append({
        "question": questions[i],
        "metrics": ragas_dict,
    })


# -------------------------------------------------
# 🔟 SAVE DOCX (simple text)
# -------------------------------------------------
doc = Document()
doc.add_heading("RAG Evaluation Report", level=1)

for i in range(len(questions)):
    q = questions[i]
    ir = ir_metrics_per_q[i]
    ragas = ragas_results_per_q[i]["metrics"]

    doc.add_heading(q, level=2)

    doc.add_paragraph("IR Metrics:")
    for k, v in ir.items():
        if k != "question":
            doc.add_paragraph(f"{k}: {v:.4f}", style="List Bullet")

    doc.add_paragraph("\nRAGAS Metrics:")
    for k, v in ragas.items():
        doc.add_paragraph(f"{k}: {v:.4f}", style="List Bullet")

doc.save("rag_evaluation_report.docx")


# -------------------------------------------------
# 1️⃣1️⃣ SAVE PDF (SINGLE-PAGE OPTIMIZED)
# -------------------------------------------------
styles = getSampleStyleSheet()

small_text = ParagraphStyle(
    "small_text", parent=styles["Normal"], fontSize=8, leading=9
)

title_style = ParagraphStyle(
    "title", parent=styles["Title"], fontSize=11, leading=12
)

flow = []

pdf = SimpleDocTemplate(
    "rag_evaluation_report.pdf",
    pagesize=A4,
    topMargin=10,
    bottomMargin=10,
    leftMargin=10,
    rightMargin=10
)

for i in range(len(questions)):
    q = questions[i]
    ir = ir_metrics_per_q[i]
    ragas = ragas_results_per_q[i]["metrics"]

    flow.append(Paragraph(f"<b><i>Question: {q}</i></b>", title_style))
    flow.append(Spacer(1, 4))

    # IR TABLE
    ir_rows = [["Metric", "Value"]]
    for k, v in ir.items():
        if k != "question":
            ir_rows.append([k, f"{v:.4f}"])

    ir_table = Table(ir_rows, colWidths=[60*mm, 25*mm])
    ir_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 7),
        ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
    ]))

    flow.append(Paragraph("<b>IR Metrics</b>", small_text))
    flow.append(ir_table)
    flow.append(Spacer(1, 4))

    # RAGAS TABLE
    ragas_rows = [["Metric", "Value"]]
    for k, v in ragas.items():
        ragas_rows.append([k, f"{v:.4f}"])

    ragas_table = Table(ragas_rows, colWidths=[60*mm, 25*mm])
    ragas_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,-1), 7),
        ('GRID', (0,0), (-1,-1), 0.25, colors.grey),
    ]))

    flow.append(Paragraph("<b>RAGAS Metrics</b>", small_text))
    flow.append(ragas_table)
    flow.append(Spacer(1, 8))

pdf.build(flow)


# -------------------------------------------------
# 1️⃣2️⃣ FINAL PRINT
# -------------------------------------------------
print("Report Saved Successfully!")
print("PDF  -> rag_evaluation_report.pdf")
print("DOCX -> rag_evaluation_report.docx")
